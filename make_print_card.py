# -*- coding: utf-8 -*-
"""יוצר קובץ Word להדפסה: כרטיסי 'פתק להקראה' בעמוד אחד, מעוצבים, מוכנים לחיתוך."""
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NOTE = (
    "שלום ג'מיני. אני אשת חינוך, וזה עתה סיימתי קורס בנוירופדגוגיה ובבינה מלאכותית.\n"
    "התפקיד שלי (מחקי את שאינו מתאים):\n"
    "מורה\n"
    "מנהלת\n"
    "סגנית\n\n"
    "אני רוצה לנהל איתך שיחה קולית רפלקטיבית על מי שאני היום, לאור הדרך שעברתי בקורס. "
    "בלב השיחה — ה-“why” שמניע אותי, ברוח “Start With Why” של סיימון סינק: "
    "לא מה אני עושה ולא איך, אלא למה.\n\n"
    "נהל אותה בנחת, שאלה אחת בכל פעם. תן לי ללכת לאן שארצה, אבל שמור אותי בתלם — "
    "בתוך הרפלקציה על הקורס ועל ה-“why” שלי — והחזר אותי בעדינות אל שאלת ה“למה”.\n\n"
    "עכשיו אני מתחילה. שאל אותי את השאלה הראשונה."
)
ACCENT = RGBColor(0x5b, 0x3e, 0x8e)
INK = RGBColor(0x2a, 0x24, 0x3a)

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)

def rtl_run(run):
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)

def set_cell_border(cell, color="C9BCE6", sz="12", dash="dashed"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), dash); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)

def shade_cell(cell, fill="FAF8FE"):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill); tcPr.append(sh)

def add_card(cell):
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    # title
    p = cell.add_paragraph(); set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(2)
    r = p.add_run("🎙️  שיחה רפלקטיבית עם הבינה"); rtl_run(r)
    r.font.name = "Heebo"; r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = ACCENT
    # body
    for block in NOTE.split("\n\n"):
        bp = cell.add_paragraph(); set_rtl(bp); bp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        bp.paragraph_format.space_after = Pt(4); bp.paragraph_format.line_spacing = 1.2
        lines = block.split("\n")
        for j, line in enumerate(lines):
            run = bp.add_run(line); rtl_run(run)
            run.font.name = "Heebo"; run.font.size = Pt(10.5); run.font.color.rgb = INK
            if j < len(lines) - 1:
                brk = bp.add_run(); brk.add_break()
    # footer tag
    fp = cell.add_paragraph(); set_rtl(fp); fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(3)
    fr = fp.add_run("✂ — — — — — — — — — — — — — — — — — — — —")
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0xB0, 0xA4, 0xCC)
    shade_cell(cell); set_cell_border(cell)

doc = Document()
sec = doc.sections[0]
sec.page_height = Mm(297); sec.page_width = Mm(210)
sec.top_margin = Mm(10); sec.bottom_margin = Mm(10)
sec.left_margin = Mm(10); sec.right_margin = Mm(10)

# heading
h = doc.add_paragraph(); set_rtl(h); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.paragraph_format.space_after = Pt(6)
hr = h.add_run("פתקי הקראה — שיחה רפלקטיבית קולית · גזרו וחלקו"); rtl_run(hr)
hr.font.name = "Heebo"; hr.font.size = Pt(12); hr.font.bold = True
hr.font.color.rgb = RGBColor(0x1f, 0x8a, 0x8a)

ROWS, COLS = 2, 2
table = doc.add_table(rows=ROWS, cols=COLS)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for row in table.rows:
    row.height = Mm(128)
    for cell in row.cells:
        cell.width = Mm(93)
        add_card(cell)

out = r"g:\האחסון שלי\תום הגלעדי - העסק\בינה מלאכותית\הרצאות קורסים וסדנאות\נוירופדגוגיה\סיכומי פגישות צוות נוירופדגוגיה\מפגש-אחרון-חומרים\פתק-הקראה-להדפסה.docx"
doc.save(out)
print("saved:", out)
print("cards:", ROWS * COLS)
